# app/utils/extract.py
import os
import mimetypes
import logging
import tempfile
import shutil
from typing import Tuple, Optional

logger = logging.getLogger(__name__)
if not logger.handlers:
    ch = logging.StreamHandler()
    ch.setFormatter(logging.Formatter("[%(levelname)s] %(asctime)s %(name)s: %(message)s"))
    logger.addHandler(ch)
logger.setLevel(logging.INFO)

# try optional libs; we'll gracefully fallback when missing
try:
    from docx import Document as DocxDocument  # python-docx
except Exception:
    DocxDocument = None

try:
    # pypdf for quick extraction
    from pypdf import PdfReader
except Exception:
    PdfReader = None

# more robust PDF extractors
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

# OCR pieces (optional)
try:
    import pytesseract
    from PIL import Image
except Exception:
    pytesseract = None
    Image = None

try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None

# cleaning helper
def _clean_whitespace(s: str) -> str:
    # collapse whitespace and strip
    return " ".join(s.split()).strip()

def _extract_text_from_docx(path: str) -> str:
    if DocxDocument is None:
        raise RuntimeError("python-docx not installed")
    doc = DocxDocument(path)
    parts = []
    # paragraphs
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    # tables
    try:
        for table in doc.tables:
            for r in table.rows:
                cells = [c.text.strip() for c in r.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception:
        # ignore table parsing errors
        logger.debug("docx: table parse error (ignored)")
    # headers & footers
    try:
        for section in doc.sections:
            hdr = section.header
            ftr = section.footer
            if hdr and hdr.paragraphs:
                for p in hdr.paragraphs:
                    if p.text and p.text.strip():
                        parts.append(p.text.strip())
            if ftr and ftr.paragraphs:
                for p in ftr.paragraphs:
                    if p.text and p.text.strip():
                        parts.append(p.text.strip())
    except Exception:
        logger.debug("docx: header/footer parse error (ignored)")
    return "\n\n".join(parts)

def _extract_text_pypdf(path: str) -> str:
    if PdfReader is None:
        raise RuntimeError("pypdf not installed")
    reader = PdfReader(path)
    pages_text = []
    for p in reader.pages:
        try:
            t = p.extract_text() or ""
        except Exception:
            t = ""
        if t and t.strip():
            pages_text.append(t)
    return "\n\n".join(pages_text)

def _extract_text_pdfplumber(path: str) -> str:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber not installed")
    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            if t and t.strip():
                texts.append(t)
    return "\n\n".join(texts)

def _extract_text_pymupdf(path: str) -> str:
    if fitz is None:
        raise RuntimeError("PyMuPDF not installed")
    doc = fitz.open(path)
    parts = []
    for page in doc:
        try:
            t = page.get_text("text") or ""
        except Exception:
            t = ""
        if t and t.strip():
            parts.append(t)
    return "\n\n".join(parts)

def _ocr_image(path_or_pil_image, lang: Optional[str] = None) -> str:
    if pytesseract is None:
        raise RuntimeError("pytesseract not installed")
    if isinstance(path_or_pil_image, str):
        img = Image.open(path_or_pil_image)
    else:
        img = path_or_pil_image
    if lang:
        return pytesseract.image_to_string(img, lang=lang)
    return pytesseract.image_to_string(img)

def _ocr_pdf(path: str, dpi: int = 300, lang: Optional[str] = None) -> str:
    if convert_from_path is None or pytesseract is None or Image is None:
        raise RuntimeError("pdf2image and/or pytesseract and PIL not installed")
    # convert each page and OCR
    texts = []
    try:
        pages = convert_from_path(path, dpi=dpi)
    except Exception as e:
        logger.exception("pdf2image.convert_from_path failed: %s", e)
        raise
    for img in pages:
        try:
            texts.append(_ocr_image(img, lang=lang))
        except Exception:
            texts.append("")
    return "\n\n".join(t for t in texts if t and t.strip())

def extract_text_from_file(path: str, original_name: str, *, ocr_threshold_chars: int = 300) -> Tuple[str, str]:
    """
    Robust text extraction for common upload types.

    Returns: (text_content, mime_type)
    - ocr_threshold_chars: if extracted text length < threshold for PDFs, attempt OCR fallback.
    """

    original_name = original_name or os.path.basename(path)
    mime_type = mimetypes.guess_type(original_name)[0] or "application/octet-stream"
    ext = os.path.splitext(original_name)[1].lower()

    logger.info("Extracting text from %s (ext=%s mime=%s)", original_name, ext, mime_type)

    # txt
    if ext in (".txt",):
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                raw = f.read()
            return _clean_whitespace(raw), mime_type
        except Exception as e:
            logger.exception("Failed to read txt file: %s", e)
            return "", mime_type

    # docx
    if ext == ".docx":
        try:
            txt = _extract_text_from_docx(path)
            txt = _clean_whitespace(txt)
            logger.info("docx extracted chars=%d", len(txt))
            return txt, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        except Exception as e:
            logger.exception("docx extraction failed: %s", e)
            # fall through to last-resort snippet

    # pdf
    if ext == ".pdf":
        # Try lightweight pypdf extraction first
        text = ""
        try:
            if PdfReader is not None:
                text = _extract_text_pypdf(path)
                logger.debug("pypdf extracted length=%d", len(text))
        except Exception as e:
            logger.debug("pypdf extraction failed: %s", e)

        # If nothing or short, try pdfplumber
        if (not text or len(text.strip()) < max(10, ocr_threshold_chars // 4)) and pdfplumber is not None:
            try:
                text2 = _extract_text_pdfplumber(path)
                if text2 and len(text2.strip()) > len(text):
                    text = text2
                logger.debug("pdfplumber extracted length=%d", len(text))
            except Exception as e:
                logger.debug("pdfplumber failed: %s", e)

        # If still short, try PyMuPDF (fitz)
        if (not text or len(text.strip()) < max(10, ocr_threshold_chars // 4)) and fitz is not None:
            try:
                text3 = _extract_text_pymupdf(path)
                if text3 and len(text3.strip()) > len(text):
                    text = text3
                logger.debug("pymupdf extracted length=%d", len(text))
            except Exception as e:
                logger.debug("pymupdf failed: %s", e)

        text = _clean_whitespace(text)
        logger.info("PDF initial extraction length=%d", len(text))

        # If text is very short, attempt OCR (scanned PDF)
        if len(text) < ocr_threshold_chars:
            logger.info("PDF looks like scanned or extracted text too small (len=%d). Trying OCR fallback.", len(text))
            try:
                ocr_text = _ocr_pdf(path)
                ocr_text = _clean_whitespace(ocr_text)
                if ocr_text and len(ocr_text) > len(text):
                    logger.info("OCR recovered text length=%d", len(ocr_text))
                    return ocr_text, "application/pdf"
                else:
                    logger.info("OCR returned insufficient text (len=%d). Using best raw extraction.", len(ocr_text))
            except Exception as e:
                logger.exception("PDF OCR fallback failed: %s", e)

        return text, "application/pdf"

    # images -> OCR
    if ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        logger.info("Image upload detected; attempting OCR.")
        try:
            if pytesseract is None or Image is None:
                raise RuntimeError("pytesseract or PIL not installed for image OCR")
            txt = _ocr_image(path)
            txt = _clean_whitespace(txt)
            logger.info("Image OCR length=%d", len(txt))
            return txt, mime_type
        except Exception as e:
            logger.exception("Image OCR failed: %s", e)
            return "", mime_type

    # last resort: return head snippet so caller has something to display/debug
    try:
        with open(path, "rb") as f:
            raw = f.read(2000)
        snippet = raw.decode("utf-8", errors="replace")
        snippet = _clean_whitespace(snippet)
        logger.warning("Falling back to binary-head snippet length=%d for %s", len(snippet), original_name)
        return snippet, mime_type
    except Exception as e:
        logger.exception("Failed to read file for fallback snippet: %s", e)
        return "", mime_type
